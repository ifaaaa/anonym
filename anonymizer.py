import re
from docx import Document

def anonymize_name(name):
    if len(name) > 1:
        return name[0] + '*' * (len(name) - 1)
    return name

def advanced_anonymizer(text):
    positions = ['шүүгч', 'прокурор', 'өмгөөлөгч', 'мөрдөн байцаагч']
    for pos in positions:
        pattern = fr'{pos}\s+[А-ЯӨҮ]\.\s?[А-ЯӨҮа-яөү]+'
        text = re.sub(pattern, f'{pos} [НЭР]', text, flags=re.IGNORECASE)

    name_pattern = r'\b[А-ЯӨҮ][а-яөү]{2,}\b'
    names = re.findall(name_pattern, text)
    for name in set(names):
        text = text.replace(name, anonymize_name(name))

    text = re.sub(r'\b[А-Яа-я]{2}\d{8}\b', '[РД]', text)
    text = re.sub(r'\b\d{8}\b', '[УТАС]', text)
    text = re.sub(r'\b[\w.-]+@[\w.-]+\.\w+\b', '[ИМЭЙЛ]', text)
    text = re.sub(r'(БЗД|СБД|ХУД|ЧД|Улаанбаатар|УБ)', '[ХАЯГ]', text)

    return text

def anonymize_docx(input_path, output_path):
    doc = Document(input_path)
    for para in doc.paragraphs:
        para.text = advanced_anonymizer(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = advanced_anonymizer(cell.text)
    doc.save(output_path)
